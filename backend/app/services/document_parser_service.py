from io import BytesIO
from pathlib import Path
import subprocess
import tempfile

import pandas as pd
from docx import Document
from pypdf import PdfReader

from app.ai.ocr_service import OCRService
from app.core.config import settings


class DocumentParserService:
    def __init__(self, ocr_service: OCRService | None = None) -> None:
        self.ocr_service = ocr_service or OCRService()

    def parse(self, filename: str, content_type: str, data: bytes | Path) -> str:
        text = self.extract_text(filename, content_type, data)
        tables = self.extract_table_text(filename, content_type, data)
        return "\n\n".join(item for item in (text, tables) if item).strip()

    def extract_text(self, filename: str, content_type: str, data: bytes | Path) -> str:
        suffix = filename.lower().split(".")[-1]
        if suffix == "pdf":
            return self._parse_pdf_text(data)
        if suffix in {"doc", "docx"}:
            return self._parse_legacy_doc(data) if suffix == "doc" else self._parse_docx_text(data)
        if suffix in {"xls", "xlsx"}:
            return self._parse_excel(data)
        if suffix == "csv":
            return self._parse_csv(data)
        if suffix in {"txt", "md"}:
            return data.decode("utf-8", errors="ignore") if isinstance(data, bytes) else data.read_text(encoding="utf-8", errors="ignore")
        if suffix in {"png", "jpg", "jpeg", "webp"}:
            image = data if isinstance(data, bytes) else data.read_bytes()
            return self.ocr_service.extract(image, content_type or f"image/{suffix}")
        raise ValueError(f"Unsupported file type: {suffix}")

    def extract_table_text(self, filename: str, content_type: str, data: bytes | Path) -> str:
        suffix = filename.lower().split(".")[-1]
        if suffix == "pdf":
            return "\n\n".join(self._extract_pdf_tables(data))
        if suffix == "docx":
            return self._extract_docx_tables(data)
        return ""

    def _parse_pdf_text(self, data: bytes | Path) -> str:
        reader = PdfReader(BytesIO(data) if isinstance(data, bytes) else str(data))
        text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        if text:
            return text
        try:
            import fitz
        except ImportError as exc:
            raise RuntimeError("PyMuPDF is required to OCR scanned PDF pages") from exc
        document = fitz.open(stream=data, filetype="pdf") if isinstance(data, bytes) else fitz.open(str(data))
        if len(document) > settings.ocr_max_pages:
            raise ValueError(f"Scanned PDF exceeds OCR page limit of {settings.ocr_max_pages}")
        pages: list[str] = []
        for page in document:
            image = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False).tobytes("png")
            pages.append(self.ocr_service.extract(image, "image/png"))
        return "\n\n".join(pages).strip()

    @staticmethod
    def _extract_pdf_tables(data: bytes | Path) -> list[str]:
        try:
            import pdfplumber
        except ImportError:
            return []
        tables: list[str] = []
        with pdfplumber.open(BytesIO(data) if isinstance(data, bytes) else str(data)) as document:
            for page_index, page in enumerate(document.pages, start=1):
                for table_index, table in enumerate(page.extract_tables(), start=1):
                    rows = [" | ".join(cell or "" for cell in row) for row in table if row]
                    if rows:
                        tables.append(f"PDF Table {page_index}.{table_index}:\n" + "\n".join(rows))
        return tables

    def _parse_docx_text(self, data: bytes | Path) -> str:
        document = Document(BytesIO(data) if isinstance(data, bytes) else str(data))
        return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())

    @staticmethod
    def _parse_legacy_doc(data: bytes | Path) -> str:
        temporary_path: Path | None = None
        try:
            if isinstance(data, bytes):
                with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as temporary:
                    temporary.write(data)
                    temporary_path = Path(temporary.name)
                source = temporary_path
            else:
                source = data
            result = subprocess.run(["antiword", str(source)], capture_output=True, text=True, timeout=60, check=False)
            if result.returncode != 0:
                raise RuntimeError(result.stderr.strip() or "Unable to parse legacy .doc file")
            return result.stdout.strip()
        except FileNotFoundError as exc:
            raise RuntimeError("antiword is required to parse legacy .doc files") from exc
        finally:
            if temporary_path:
                temporary_path.unlink(missing_ok=True)

    def _extract_docx_tables(self, data: bytes | Path) -> str:
        document = Document(BytesIO(data) if isinstance(data, bytes) else str(data))
        sections: list[str] = []
        for table_index, table in enumerate(document.tables, start=1):
            sections.append(f"Table {table_index}:")
            for row in table.rows:
                sections.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(sections)

    def _parse_excel(self, data: bytes | Path) -> str:
        workbook = pd.read_excel(BytesIO(data) if isinstance(data, bytes) else data, sheet_name=None)
        sections: list[str] = []
        for sheet_name, frame in workbook.items():
            sections.append(f"Sheet: {sheet_name}")
            sections.append(frame.fillna("").to_csv(index=False))
        return "\n".join(sections)

    def _parse_csv(self, data: bytes | Path) -> str:
        frame = pd.read_csv(BytesIO(data) if isinstance(data, bytes) else data)
        return frame.fillna("").to_csv(index=False)
