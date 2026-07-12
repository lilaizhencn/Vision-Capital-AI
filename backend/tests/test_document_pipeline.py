from io import BytesIO
import hashlib
import socket
from datetime import datetime, timedelta, timezone

import pandas as pd
import pytest
import fitz
from PIL import Image
from types import SimpleNamespace
from docx import Document
from app.ai.llm_service import LLMService
from app.ai.embedding_service import EmbeddingService
from app.services.rag_service import RAGService
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app.rag.chunking import chunk_text, estimate_tokens
from app.services.document_parser_service import DocumentParserService
from app.main import app
from app.core.database import Base, get_db
from app.storage.storage_service import LocalStorageService
from app.core.config import Settings, settings
from app.workers import tasks
from app.services import file_service
from app.api import websocket as websocket_api
from app.services.file_service import FileService
from app.services.file_validation_service import validate_file_signature
from app.services.virus_scan_service import VirusScanner


class FakeOCR:
    def extract(self, data: bytes, content_type: str) -> str:
        assert data
        assert content_type.startswith("image/")
        return "OCR result"


def test_parser_extracts_docx_text() -> None:
    document = Document()
    document.add_paragraph("Investment thesis")
    stream = BytesIO()
    document.save(stream)

    result = DocumentParserService().parse("memo.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", stream.getvalue())

    assert "Investment thesis" in result


def test_parser_extracts_docx_tables() -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "100"
    stream = BytesIO()
    document.save(stream)

    result = DocumentParserService().parse("table.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", stream.getvalue())

    assert "Table 1:" in result
    assert "Revenue | 100" in result


def test_llm_extraction_parses_json(monkeypatch) -> None:
    service = LLMService()
    monkeypatch.setattr(service, "generate", lambda _prompt: "```json\n{\"risks\": [\"liquidity\"]}\n```")

    assert service.extract_document_data("document") == {"risks": ["liquidity"]}


def test_parser_extracts_excel_sheets() -> None:
    stream = BytesIO()
    with pd.ExcelWriter(stream, engine="openpyxl") as writer:
        pd.DataFrame({"metric": ["revenue"], "value": [100]}).to_excel(writer, index=False, sheet_name="Metrics")

    result = DocumentParserService().parse("metrics.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", stream.getvalue())

    assert "Sheet: Metrics" in result
    assert "revenue" in result


def test_parser_extracts_pdf_tables() -> None:
    document = fitz.open()
    page = document.new_page(width=300, height=200)
    for y in (20, 60, 100):
        page.draw_line(fitz.Point(20, y), fitz.Point(260, y), color=(0, 0, 0), width=1)
    for x in (20, 140, 260):
        page.draw_line(fitz.Point(x, 20), fitz.Point(x, 100), color=(0, 0, 0), width=1)
    page.insert_text((30, 45), "Metric")
    page.insert_text((150, 45), "Value")
    page.insert_text((30, 85), "Revenue")
    page.insert_text((150, 85), "100")

    result = DocumentParserService().parse("table.pdf", "application/pdf", document.tobytes())

    assert "PDF Table 1.1:" in result
    assert "Revenue | 100" in result


def test_large_pdf_table_extraction_samples_bounded_pages(monkeypatch) -> None:
    import sys

    visited: list[int] = []

    class FakePage:
        def __init__(self, index: int):
            self.index = index

        def extract_tables(self):
            visited.append(self.index)
            return []

    class FakeDocument:
        pages = [FakePage(index) for index in range(200)]

        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

    monkeypatch.setattr(settings, "pdf_table_max_pages", 10)
    monkeypatch.setitem(sys.modules, "pdfplumber", SimpleNamespace(open=lambda _source: FakeDocument()))

    assert DocumentParserService._extract_pdf_tables(b"%PDF") == []
    assert len(visited) == 10
    assert visited[0] == 0
    assert visited[-1] == 199


def test_parser_delegates_images_to_ocr() -> None:
    result = DocumentParserService(ocr_service=FakeOCR()).parse("scan.png", "image/png", b"image-bytes")

    assert result == "OCR result"


@pytest.mark.parametrize(
    ("filename", "content_type", "image_format"),
    [("scan.jpg", "image/jpeg", "JPEG"), ("scan.webp", "image/webp", "WEBP")],
)
def test_parser_delegates_jpeg_and_webp_to_ocr(filename: str, content_type: str, image_format: str) -> None:
    stream = BytesIO()
    Image.new("RGB", (32, 32), "white").save(stream, format=image_format)

    result = DocumentParserService(ocr_service=FakeOCR()).parse(filename, content_type, stream.getvalue())

    assert result == "OCR result"


def test_scanned_pdf_uses_ocr_when_no_text_layer() -> None:
    document = fitz.open()
    document.new_page(width=100, height=100)

    result = DocumentParserService(ocr_service=FakeOCR()).parse("scan.pdf", "application/pdf", document.tobytes())

    assert result == "OCR result"


def test_ocr_falls_back_to_local_tesseract_when_vision_provider_fails(monkeypatch) -> None:
    from app.ai.ocr_service import OCRService

    service = OCRService()
    service.client = SimpleNamespace(chat=SimpleNamespace(completions=SimpleNamespace(create=lambda **_kwargs: (_ for _ in ()).throw(ValueError("vision unsupported")))))
    monkeypatch.setattr(service, "_extract_local", lambda _data: "local OCR result")

    assert service.extract(b"image-bytes", "image/png") == "local OCR result"


def test_file_validation_rejects_unsupported_extensions() -> None:
    with pytest.raises(Exception):
        FileService._validate_filename("payload.exe")

    with pytest.raises(Exception):
        FileService._validate_filename("no-extension")


def test_file_signature_validation_rejects_extension_spoofing() -> None:
    with pytest.raises(ValueError, match="signature"):
        validate_file_signature("fake.pdf", b"not a pdf")
    validate_file_signature("notes.txt", b"plain text")


def test_virus_scanner_is_explicitly_skipped_only_outside_production(monkeypatch) -> None:
    monkeypatch.setattr(settings, "virus_scan_enabled", False)
    assert VirusScanner().scan_bytes(b"safe") == "skipped"


def test_production_configuration_requires_llm_credentials() -> None:
    production = Settings(
        app_env="production",
        app_secret_key="app-secret",
        jwt_secret_key="jwt-secret",
        database_url="postgresql+psycopg://user:password@db/vision",
        r2_endpoint_url="https://example.r2.cloudflarestorage.com",
        r2_access_key_id="access",
        r2_secret_access_key="secret",
        r2_bucket_name="app-public",
        virus_scan_enabled=True,
    )

    with pytest.raises(RuntimeError, match="LLM_API_KEY"):
        production.validate_production()

    production.llm_api_key = "test-key"
    production.validate_production()


def test_virus_scanner_streams_clamav_protocol(monkeypatch) -> None:
    class FakeConnection:
        def __init__(self):
            self.payload = bytearray()

        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

        def sendall(self, value):
            self.payload.extend(value)

        def recv(self, _size):
            return b"stream: OK\x00"

    fake = FakeConnection()
    monkeypatch.setattr(settings, "virus_scan_enabled", True)
    monkeypatch.setattr(socket, "create_connection", lambda *_args, **_kwargs: fake)
    assert VirusScanner().scan_bytes(b"safe") == "stream: OK"
    assert bytes(fake.payload).startswith(b"zINSTREAM\x00")


def test_virus_scanner_unavailability_is_retryable(monkeypatch) -> None:
    from app.services.virus_scan_service import VirusScannerUnavailable

    monkeypatch.setattr(settings, "virus_scan_enabled", True)
    monkeypatch.setattr(socket, "create_connection", lambda *_args, **_kwargs: (_ for _ in ()).throw(OSError("busy")))
    with pytest.raises(VirusScannerUnavailable):
        VirusScanner().scan_bytes(b"safe")
    assert issubclass(VirusScannerUnavailable, ConnectionError)


def test_chunking_overlaps_and_estimates_tokens() -> None:
    text = "0123456789" * 300

    chunks = chunk_text(text, chunk_size=100, overlap=20)

    assert len(chunks) > 1
    assert chunks[0][-20:] == chunks[1][:20]
    assert estimate_tokens(text) > 0


def test_parse_progress_matches_solution_stage_weights() -> None:
    assert [tasks.STAGE_PROGRESS[stage] for stage in tasks.STAGE_PROGRESS] == [10, 20, 60, 80, 95]


def test_local_storage_round_trip(tmp_path, monkeypatch) -> None:
    monkeypatch.setattr(settings, "local_storage_path", tmp_path)
    storage = LocalStorageService()

    storage.upload_file("batch/file.txt", b"hello", "text/plain")

    assert storage.object_exists("batch/file.txt", expected_size=5)
    assert storage.download_file("batch/file.txt") == b"hello"
    destination = tmp_path / "streamed" / "file.txt"
    storage.download_file_to_path("batch/file.txt", destination)
    assert destination.read_bytes() == b"hello"


def test_health_contract_without_infrastructure() -> None:
    client = TestClient(app)

    assert client.get("/health").json() == {"status": "ok"}
    assert client.get("/health/ready").status_code == 503


@pytest.fixture
def api_client(tmp_path, monkeypatch):
    engine = create_engine("sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool)
    Base.metadata.create_all(engine)
    TestSession = sessionmaker(bind=engine, autoflush=False, autocommit=False, expire_on_commit=False)

    def override_get_db():
        db = TestSession()
        try:
            yield db
        finally:
            db.close()

    app.dependency_overrides[get_db] = override_get_db
    monkeypatch.setattr(tasks, "SessionLocal", TestSession)
    monkeypatch.setattr(websocket_api, "SessionLocal", TestSession)
    monkeypatch.setattr(settings, "local_storage_path", tmp_path)
    monkeypatch.setattr(settings, "celery_task_always_eager", True)
    monkeypatch.setattr(settings, "research_auto_enrich_enabled", False)
    eager_parse = lambda file_id: tasks.parse_uploaded_file_task.run(file_id)
    monkeypatch.setattr(file_service, "parse_uploaded_file_task", SimpleNamespace(delay=eager_parse))
    client = TestClient(app)
    yield client
    app.dependency_overrides.clear()


def test_auth_project_upload_and_parse_flow(api_client) -> None:
    registered = api_client.post("/api/auth/register", json={"email": "qa@example.com", "username": "qa", "password": "strong-password"})
    assert registered.status_code == 200
    token = registered.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    logged_in = api_client.post("/api/auth/login", json={"email": "qa@example.com", "password": "strong-password"})
    assert logged_in.status_code == 200
    assert api_client.get("/api/auth/me", headers=headers).json()["username"] == "qa"

    project = api_client.post("/api/projects", headers=headers, json={"name": "QA", "company_name": "Acme", "industry": "SaaS", "stage": "Seed"})
    assert project.status_code == 200
    project_id = project.json()["id"]

    uploaded = api_client.post(
        f"/api/projects/{project_id}/files/upload",
        headers=headers,
        files={"upload_file": ("memo.txt", b"Revenue grew 20 percent", "text/plain")},
    )
    assert uploaded.status_code == 200

    listed = api_client.get(f"/api/projects/{project_id}/files", headers=headers)
    assert listed.status_code == 200
    assert listed.json()[0]["filename"] == "memo.txt"
    assert listed.json()[0]["parse_status"] == "completed"
    assert listed.json()[0]["progress"] == 100
    downloaded = api_client.get(f"/api/files/{listed.json()[0]['id']}/download", headers=headers)
    assert downloaded.status_code == 200
    assert downloaded.content == b"Revenue grew 20 percent"
    assert "attachment" in downloaded.headers["content-disposition"]


def test_single_upload_enforces_maximum_size(api_client, monkeypatch) -> None:
    token = api_client.post("/api/auth/register", json={"email": "size@example.com", "username": "size", "password": "strong-password"}).json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    project_id = api_client.post("/api/projects", headers=headers, json={"name": "Size QA", "company_name": "Acme", "industry": "SaaS", "stage": "Seed"}).json()["id"]
    monkeypatch.setattr(settings, "max_upload_size_bytes", 3)

    response = api_client.post(
        f"/api/projects/{project_id}/files/upload",
        headers=headers,
        files={"upload_file": ("memo.txt", b"four", "text/plain")},
    )

    assert response.status_code == 413


def test_duplicate_clean_file_reuses_existing_parse(api_client, monkeypatch) -> None:
    monkeypatch.setattr(VirusScanner, "scan_file", lambda _self, _path: "stream: OK")
    token = api_client.post(
        "/api/auth/register",
        json={"email": "dedupe@example.com", "username": "dedupe", "password": "strong-password"},
    ).json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    project_id = api_client.post(
        "/api/projects", headers=headers,
        json={"name": "Dedupe QA", "company_name": "Acme", "industry": "SaaS", "stage": "Seed"},
    ).json()["id"]
    for filename in ("original.txt", "duplicate.txt"):
        response = api_client.post(
            f"/api/projects/{project_id}/files/upload", headers=headers,
            files={"upload_file": (filename, b"same clean investment evidence", "text/plain")},
        )
        assert response.status_code == 200

    files = api_client.get(f"/api/projects/{project_id}/files", headers=headers).json()
    duplicate = next(item for item in files if item["filename"] == "duplicate.txt")
    original = next(item for item in files if item["filename"] == "original.txt")
    assert duplicate["parse_status"] == "completed"
    assert duplicate["extracted_data"] == {"duplicate_of": original["id"]}


def test_batch_upload_complete_and_parse_flow(api_client) -> None:
    token = api_client.post("/api/auth/register", json={"email": "batch@example.com", "username": "batch", "password": "strong-password"}).json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    project_id = api_client.post("/api/projects", headers=headers, json={"name": "Batch QA", "company_name": "Acme", "industry": "SaaS", "stage": "Seed"}).json()["id"]

    created = api_client.post(f"/api/projects/{project_id}/file-batches", headers=headers, json={"files": [{"filename": "batch.txt", "size": 5, "content_type": "text/plain", "checksum_sha256": hashlib.sha256(b"hello").hexdigest()}]})
    assert created.status_code == 200
    batch = created.json()
    file_id = batch["files"][0]["id"]
    resumable = api_client.get(f"/api/file-batches/{batch['id']}", headers=headers)
    assert resumable.status_code == 200
    assert resumable.json()["upload_sessions"][0]["file_id"] == file_id

    uploaded = api_client.post(f"/api/file-batches/{batch['id']}/files/{file_id}/content", headers=headers, files={"upload_file": ("batch.txt", b"hello", "text/plain")})
    assert uploaded.status_code == 200
    completed = api_client.post(f"/api/file-batches/{batch['id']}/complete", headers=headers)
    assert completed.status_code == 200

    with api_client.websocket_connect(f"/api/ws/batches/{batch['id']}?token={token}") as websocket:
        progress = websocket.receive_json()
        assert progress["batch_id"] == batch["id"]
        assert progress["status"] == "completed"

    files = api_client.get(f"/api/projects/{project_id}/files", headers=headers).json()
    assert files[0]["parse_status"] == "completed"
    assert files[0]["batch_id"] == batch["id"]
    assert files[0]["checksum_sha256"] == hashlib.sha256(b"hello").hexdigest()
    assert files[0]["virus_scan_status"] == "skipped"
    repeated = api_client.post(f"/api/file-batches/{batch['id']}/complete", headers=headers)
    assert repeated.status_code == 200
    assert repeated.json()["status"] == "completed"


def test_batch_checksum_mismatch_fails_parse(api_client) -> None:
    token = api_client.post("/api/auth/register", json={"email": "checksum@example.com", "username": "checksum", "password": "strong-password"}).json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    project_id = api_client.post(f"/api/projects", headers=headers, json={"name": "Checksum QA", "company_name": "Acme", "industry": "SaaS", "stage": "Seed"}).json()["id"]
    created = api_client.post(
        f"/api/projects/{project_id}/file-batches",
        headers=headers,
        json={"files": [{"filename": "checksum.txt", "size": 5, "content_type": "text/plain", "checksum_sha256": "0" * 64}]},
    ).json()
    file_id = created["files"][0]["id"]
    api_client.post(f"/api/file-batches/{created['id']}/files/{file_id}/content", headers=headers, files={"upload_file": ("checksum.txt", b"hello", "text/plain")})
    with pytest.raises(ValueError, match="checksum"):
        api_client.post(f"/api/file-batches/{created['id']}/complete", headers=headers)
    file_item = api_client.get(f"/api/projects/{project_id}/files", headers=headers).json()[0]
    assert file_item["parse_status"] == "failed"
    assert "checksum" in file_item["parse_error"].lower()


def test_project_dashboard_chat_and_report_flow(api_client, monkeypatch) -> None:
    monkeypatch.setattr(RAGService, "similarity_search", lambda _self, *_args, **_kwargs: [])
    monkeypatch.setattr(LLMService, "generate", lambda _self, _prompt: "stub AI response")
    token = api_client.post("/api/auth/register", json={"email": "business@example.com", "username": "business", "password": "strong-password"}).json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}

    created = api_client.post(
        "/api/projects",
        headers=headers,
        json={"name": "Business QA", "company_name": "Acme", "industry": "SaaS", "stage": "Seed"},
    )
    assert created.status_code == 200
    project_id = created.json()["id"]

    updated = api_client.put(
        f"/api/projects/{project_id}",
        headers=headers,
        json={"name": "Business QA Updated", "company_name": "Acme", "industry": "SaaS", "stage": "Series A", "description": "Updated"},
    )
    assert updated.status_code == 200
    assert updated.json()["stage"] == "Series A"
    assert api_client.get("/api/projects", headers=headers).json()[0]["name"] == "Business QA Updated"

    dashboard = api_client.get("/api/dashboard/summary", headers=headers)
    assert dashboard.status_code == 200
    assert dashboard.json()["total_projects"] == 1

    chat = api_client.post(f"/api/projects/{project_id}/chat", headers=headers, json={"message": "Summarize the project"})
    assert chat.status_code == 200
    assert chat.json()["answer"] == "stub AI response"
    assert chat.json()["citations"] == []
    assert chat.json()["confidence"] == "low"
    assert chat.json()["missing_evidence"]

    generated = api_client.post(f"/api/projects/{project_id}/reports/generate", headers=headers)
    assert generated.status_code == 200
    assert generated.json()["content"] == "stub AI response"
    listed_reports = api_client.get(f"/api/projects/{project_id}/reports", headers=headers)
    assert listed_reports.status_code == 200
    assert len(listed_reports.json()) == 1
    recent_reports = api_client.get("/api/reports", headers=headers)
    assert recent_reports.status_code == 200
    assert len(recent_reports.json()) == 1

    other_token = api_client.post(
        "/api/auth/register",
        json={"email": "other-business@example.com", "username": "other-business", "password": "strong-password"},
    ).json()["access_token"]
    other_headers = {"Authorization": f"Bearer {other_token}"}
    assert api_client.get("/api/reports", headers=other_headers).json() == []
    assert api_client.get(f"/api/projects/{project_id}/reports", headers=other_headers).status_code == 404

    assert api_client.delete(f"/api/projects/{project_id}", headers=headers).status_code == 200
    assert api_client.get(f"/api/projects/{project_id}", headers=headers).status_code == 404


def test_cross_tenant_project_resources_and_storage_are_isolated(api_client, monkeypatch) -> None:
    monkeypatch.setattr(LLMService, "generate", lambda _self, _prompt: "isolated")
    owner_token = api_client.post(
        "/api/auth/register",
        json={"email": "tenant-owner@example.com", "username": "tenant-owner", "password": "strong-password"},
    ).json()["access_token"]
    other_token = api_client.post(
        "/api/auth/register",
        json={"email": "tenant-other@example.com", "username": "tenant-other", "password": "strong-password"},
    ).json()["access_token"]
    owner_headers = {"Authorization": f"Bearer {owner_token}"}
    other_headers = {"Authorization": f"Bearer {other_token}"}
    owner_id = api_client.get("/api/auth/me", headers=owner_headers).json()["id"]
    project_id = api_client.post(
        "/api/projects", headers=owner_headers,
        json={"name": "Tenant Boundary", "company_name": "Boundary Co", "industry": "Security", "stage": "Seed"},
    ).json()["id"]
    uploaded = api_client.post(
        f"/api/projects/{project_id}/files/upload", headers=owner_headers,
        files={"upload_file": ("tenant.txt", b"confidential customer contract", "text/plain")},
    ).json()["file"]
    batch = api_client.post(
        f"/api/projects/{project_id}/file-batches", headers=owner_headers,
        json={"files": [{"filename": "private.txt", "size": 7, "content_type": "text/plain"}]},
    ).json()

    assert uploaded["r2_object_key"].startswith(f"tenants/{owner_id}/{project_id}/")
    assert api_client.get(f"/api/projects/{project_id}", headers=other_headers).status_code == 404
    assert api_client.get(f"/api/projects/{project_id}/files", headers=other_headers).status_code == 404
    assert api_client.get(f"/api/files/{uploaded['id']}", headers=other_headers).status_code == 404
    assert api_client.get(f"/api/file-batches/{batch['id']}", headers=other_headers).status_code == 404
    assert api_client.post(f"/api/file-batches/{batch['id']}/complete", headers=other_headers).status_code == 404
    assert api_client.get(f"/api/projects/{project_id}/research", headers=other_headers).status_code == 404
    assert api_client.post(f"/api/projects/{project_id}/chat", headers=other_headers, json={"message": "show data"}).status_code == 404
    assert api_client.get(f"/api/projects/{project_id}/tasks", headers=other_headers).status_code == 404
    assert api_client.get(f"/api/projects/{project_id}/monitoring", headers=other_headers).status_code == 404

    workspace = api_client.get(f"/api/projects/{project_id}/research", headers=owner_headers)
    assert workspace.status_code == 200
    assert len(workspace.json()["requirements"]) == 8


def test_research_trust_allowlist_rejects_lookalike_domains() -> None:
    from app.services.research_service import ResearchService

    assert ResearchService._is_trusted_domain("www.sec.gov")
    assert ResearchService._is_trusted_domain("documents1.worldbank.org")
    assert ResearchService._is_trusted_domain("www.irena.org")
    assert ResearchService._is_trusted_domain("www.federalreserve.gov")
    assert ResearchService._is_trusted_domain("tsapps.nist.gov")
    assert ResearchService._is_trusted_domain("www.census.gov")
    assert not ResearchService._is_trusted_domain("sec.gov.attacker.example")
    assert not ResearchService._is_trusted_domain("example.com")


def test_research_relevance_rejects_authoritative_but_unrelated_results() -> None:
    from app.services.research_service import ResearchService

    project = SimpleNamespace(company_name="Moderna, Inc.", industry="Biotechnology")
    assert ResearchService._content_is_relevant(project, "financial", "Moderna annual report. Moderna reported cash flow.")
    assert not ResearchService._content_is_relevant(project, "financial", "Another issuer annual report")
    assert ResearchService._content_is_relevant(project, "market", "Global biotechnology sector outlook")
    assert not ResearchService._content_is_relevant(project, "market", "Maritime transport industry outlook")
    broad_project = SimpleNamespace(company_name="Snowflake Inc.", industry="Enterprise software and AI data cloud")
    assert not ResearchService._content_is_relevant(broad_project, "market", "Enterprise survey overview")
    assert ResearchService._content_is_relevant(broad_project, "market", "Enterprise software and cloud market outlook")


def test_sec_research_rejects_stale_and_valuation_filings() -> None:
    from app.services.research_service import ResearchService

    year = datetime.now(timezone.utc).year
    current = f"Walmart Inc. Annual Report on Form 10-K for {year}"
    assert ResearchService._sec_document_is_current_filing("business", f"https://www.sec.gov/{year}/report.htm", current)
    assert not ResearchService._sec_document_is_current_filing("business", "https://www.sec.gov/2017/report.htm", "Form 10-K 2017")
    assert not ResearchService._sec_document_is_current_filing("valuation", f"https://www.sec.gov/{year}/424b2.htm", current)


def test_rag_company_match_normalizes_legal_suffixes() -> None:
    assert RAGService._same_company("JPMorgan Chase & Co.", "JPMorganChase")
    assert RAGService._same_company("Caterpillar Inc.", "Caterpillar Incorporated")
    assert not RAGService._same_company("Caterpillar Inc.", "Royal Bank of Canada")


def test_monitoring_updates_are_persisted_and_owner_scoped(api_client) -> None:
    token = api_client.post("/api/auth/register", json={"email": "monitor@example.com", "username": "monitor", "password": "strong-password"}).json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    project_id = api_client.post(
        "/api/projects",
        headers=headers,
        json={"name": "Monitor QA", "company_name": "Acme", "industry": "SaaS", "stage": "Series A"},
    ).json()["id"]

    created = api_client.post(
        f"/api/projects/{project_id}/monitoring",
        headers=headers,
        json={"metric_name": "月度收入", "metric_value": "320", "metric_unit": "万元", "risk_level": "watch", "note": "客户集中度需要继续跟踪"},
    )
    assert created.status_code == 200
    assert created.json()["risk_level"] == "watch"
    listed = api_client.get(f"/api/projects/{project_id}/monitoring", headers=headers)
    assert listed.status_code == 200
    assert listed.json()[0]["metric_name"] == "月度收入"

    other_token = api_client.post("/api/auth/register", json={"email": "other-monitor@example.com", "username": "other-monitor", "password": "strong-password"}).json()["access_token"]
    assert api_client.get(f"/api/projects/{project_id}/monitoring", headers={"Authorization": f"Bearer {other_token}"}).status_code == 404


def test_project_tasks_are_persisted_and_owner_scoped(api_client) -> None:
    token = api_client.post("/api/auth/register", json={"email": "tasks@example.com", "username": "tasks", "password": "strong-password"}).json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    project_id = api_client.post(
        "/api/projects", headers=headers,
        json={"name": "Tasks QA", "company_name": "Acme", "industry": "SaaS", "stage": "Seed"},
    ).json()["id"]
    listed = api_client.get(f"/api/projects/{project_id}/tasks", headers=headers)
    assert listed.status_code == 200
    assert len(listed.json()) == 3
    task = listed.json()[0]
    rejected = api_client.patch(f"/api/projects/{project_id}/tasks/{task['id']}", headers=headers, json={"done": True})
    assert rejected.status_code == 400
    updated = api_client.patch(
        f"/api/projects/{project_id}/tasks/{task['id']}",
        headers=headers,
        json={"status": "completed", "result": "Management biographies were reconciled to the proxy filing."},
    )
    assert updated.status_code == 200
    assert updated.json()["done"] is True
    assert updated.json()["status"] == "completed"
    assert updated.json()["completed_at"] is not None


def test_requirement_detail_exposes_fields_evidence_and_task_link(api_client) -> None:
    token = api_client.post(
        "/api/auth/register",
        json={"email": "requirement@example.com", "username": "requirement", "password": "strong-password"},
    ).json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    project_id = api_client.post(
        "/api/projects",
        headers=headers,
        json={"name": "Requirement QA", "company_name": "Acme", "industry": "SaaS", "stage": "Seed"},
    ).json()["id"]
    uploaded = api_client.post(
        f"/api/projects/{project_id}/files/upload",
        headers=headers,
        files={"upload_file": (
            "team.txt",
            b"Chief Executive Officer Jane Doe has 20 years of management experience. The Board of Directors includes independent directors.",
            "text/plain",
        )},
    )
    file_id = uploaded.json()["file"]["id"]
    workspace = api_client.get(f"/api/projects/{project_id}/research", headers=headers).json()
    team = next(item for item in workspace["requirements"] if item["category"] == "team")

    detail = api_client.get(
        f"/api/projects/{project_id}/research/requirements/{team['id']}", headers=headers
    )
    assert detail.status_code == 200
    assert any(field["status"] == "found" for field in detail.json()["fields"])
    assert detail.json()["related_files"][0]["id"] == file_id

    rejected = api_client.post(
        f"/api/projects/{project_id}/tasks",
        headers=headers,
        json={"label": "Incomplete QA", "status": "completed"},
    )
    assert rejected.status_code == 400

    created = api_client.post(
        f"/api/projects/{project_id}/tasks",
        headers=headers,
        json={
            "label": "核验核心管理层履历",
            "description": "逐项核验管理层任职经历与董事会独立性。",
            "assignee": "Investment research",
            "due_date": "2026-07-31",
            "status": "completed",
            "result": "Management biographies were checked against the filing.",
            "evidence_file_ids": [file_id],
            "related_requirement_id": team["id"],
        },
    )
    assert created.status_code == 201
    assert created.json()["related_requirement_id"] == team["id"]
    assert created.json()["status"] == "completed"
    assert created.json()["done"] is True
    assert created.json()["evidence_file_ids"] == [file_id]
    assert created.json()["completed_at"] is not None


def test_requirement_detail_rejects_filing_table_of_contents_as_evidence() -> None:
    from app.services.research_service import ResearchService

    content = (
        "Part III Item 10 Directors, Executive Officers and Corporate Governance 82 "
        "Item 11 Executive Compensation 82 Item 12 Security Ownership of Certain Beneficial Owners 82"
    )
    assert ResearchService._is_substantive_evidence(content) is False
    assert ResearchService._is_substantive_evidence(
        "The compensation committee approved a three-year long-term incentive award for the executive team."
    ) is True


def test_chat_falls_back_to_recent_chunks_when_embeddings_are_unavailable(api_client, monkeypatch) -> None:
    monkeypatch.setattr(EmbeddingService, "embed_text", lambda _self, _text: (_ for _ in ()).throw(RuntimeError("embedding unavailable")))
    monkeypatch.setattr(LLMService, "generate", lambda _self, _prompt: "fallback response")
    token = api_client.post("/api/auth/register", json={"email": "rag@example.com", "username": "rag", "password": "strong-password"}).json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    project_id = api_client.post(
        "/api/projects",
        headers=headers,
        json={"name": "RAG QA", "company_name": "Acme", "industry": "SaaS", "stage": "Seed"},
    ).json()["id"]
    uploaded = api_client.post(
        f"/api/projects/{project_id}/files/upload",
        headers=headers,
        files={"upload_file": ("notes.txt", b"Revenue grew 20 percent", "text/plain")},
    )
    assert uploaded.status_code == 200

    response = api_client.post(f"/api/projects/{project_id}/chat", headers=headers, json={"message": "What grew?"})

    assert response.status_code == 200
    assert response.json()["answer"] == "fallback response"
    assert response.json()["citations"][0]["filename"] == "notes.txt"


def test_persist_stage_tolerates_embedding_provider_errors(api_client, monkeypatch) -> None:
    monkeypatch.setattr(EmbeddingService, "embed_text", lambda _self, _text: (_ for _ in ()).throw(ValueError("embedding endpoint returned 404")))
    token = api_client.post("/api/auth/register", json={"email": "embedding-qa@example.com", "username": "embeddingqa", "password": "strong-password"}).json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    project_id = api_client.post("/api/projects", headers=headers, json={"name": "Embedding QA", "company_name": "Acme", "industry": "SaaS", "stage": "Seed"}).json()["id"]

    uploaded = api_client.post(
        f"/api/projects/{project_id}/files/upload",
        headers=headers,
        files={"upload_file": ("notes.txt", b"Revenue grew 20 percent", "text/plain")},
    )

    assert uploaded.status_code == 200
    file_item = api_client.get(f"/api/projects/{project_id}/files", headers=headers).json()[0]
    assert file_item["parse_status"] == "completed"
    assert file_item["progress"] == 100


def test_stale_parse_stage_is_automatically_requeued(monkeypatch) -> None:
    from app.models.file import ParseStage, ParseStageRun, ParseStatus, ProjectFile

    engine = create_engine("sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool)
    Base.metadata.create_all(engine)
    TestSession = sessionmaker(bind=engine, autoflush=False, autocommit=False, expire_on_commit=False)
    with TestSession() as db:
        file = ProjectFile(
            project_id="project-1",
            filename="interrupted.pdf",
            content_type="application/pdf",
            size=100,
            r2_object_key="tenants/owner/project/interrupted.pdf",
            parse_status=ParseStatus.processing,
            parse_stage=ParseStage.ocr,
            progress=20,
        )
        db.add(file)
        db.flush()
        db.add(ParseStageRun(
            file_id=file.id,
            stage=ParseStage.ocr.value,
            idempotency_key=f"single:{file.id}:ocr",
            status="running",
            attempts=1,
            started_at=datetime.now(timezone.utc) - timedelta(minutes=31),
        ))
        db.commit()
        file_id = file.id

    requeued: list[str] = []
    monkeypatch.setattr(tasks, "SessionLocal", TestSession)
    monkeypatch.setattr(settings, "parse_stale_after_minutes", 30)
    monkeypatch.setattr(tasks.parse_uploaded_file_task, "delay", lambda value: requeued.append(value))

    assert tasks.recover_stale_parse_tasks_task.run() == {"recovered": 1}
    assert requeued == [file_id]
    with TestSession() as db:
        recovered = db.get(ProjectFile, file_id)
        assert recovered.parse_status == ParseStatus.pending
        assert "Recovering" in recovered.parse_error


def test_llm_service_uses_openai_compatible_chat_completions(monkeypatch) -> None:
    service = LLMService()
    calls = {}

    class FakeCompletions:
        def create(self, **kwargs):
            calls.update(kwargs)
            return SimpleNamespace(choices=[SimpleNamespace(message=SimpleNamespace(content="answer"))])

    service.client = SimpleNamespace(chat=SimpleNamespace(completions=FakeCompletions()))

    assert service.generate("test prompt") == "answer"
    assert calls["messages"] == [{"role": "user", "content": "test prompt"}]


def test_chinese_investment_stage_questions_use_strategy_path() -> None:
    from app.services.chat_service import ChatService

    assert ChatService._is_strategy_question("请给出投前策略")
    assert ChatService._is_strategy_question("投中阶段如何控制仓位")
    assert ChatService._is_strategy_question("投后应该跟踪哪些 KPI")
    assert not ChatService._is_strategy_question("请概括产品功能")


def test_strategy_numeric_guard_removes_values_absent_from_evidence() -> None:
    from app.services.chat_service import ChatService

    answer = "## 1. 事实\n收入为 120 百万美元。\n建议设置 65% 毛利率门槛。\n保持定性观察。"
    guarded = ChatService._remove_unsupported_numeric_lines(answer, ["Revenue was $120 million."])

    assert "120" in guarded
    assert "65%" not in guarded
    assert "保持定性观察" in guarded


def test_strategy_guard_removes_invented_consecutive_period_exit_rules() -> None:
    from app.services.chat_service import ChatService

    answer = "保留证据核验动作。连续两个季度未达预期则触发退出。继续跟踪原始披露。"
    guarded = ChatService._remove_unsupported_numeric_lines(answer, ["Company reported quarterly results."])

    assert "连续两个季度" not in guarded
    assert "保留证据核验动作" in guarded
    assert "继续跟踪原始披露" in guarded


def test_evidence_ledger_uses_exact_quotes_and_role_ids() -> None:
    from app.schemas.chat import Citation
    from app.services.evidence_ledger_service import EvidenceLedgerService

    citations = [
        Citation(
            file_id="company-file",
            filename="annual-report.pdf",
            content="The company reported revenue of $120 million and positive operating cash flow for the fiscal year.",
            document_role="company_disclosure",
        ),
        Citation(
            file_id="industry-file",
            filename="industry-report.pdf",
            content="The manufacturing industry experienced broad market uncertainty and supply-chain risk during the year.",
            document_role="industry_context",
        ),
    ]

    ledger = EvidenceLedgerService.build(citations)

    assert {item.claim_id[0] for item in ledger} == {"C", "I"}
    assert all(item.claim == item.evidence_quote for item in ledger)
    assert all(item.evidence_quote in next(c.content for c in citations if c.filename == item.source_filename) for item in ledger)


def test_evidence_ledger_rejects_unknown_and_unreferenced_company_claims() -> None:
    from app.schemas.chat import Citation
    from app.services.evidence_ledger_service import EvidenceLedgerService

    ledger = EvidenceLedgerService.build([
        Citation(
            file_id="company-file",
            filename="annual-report.pdf",
            content="The company reported revenue of $120 million and positive operating cash flow for the fiscal year.",
            document_role="company_disclosure",
        )
    ])

    missing = EvidenceLedgerService.reference_issues("Company-disclosed facts: Revenue was reported.", ledger)
    unknown = EvidenceLedgerService.reference_issues("Company-disclosed facts: Revenue was reported [C99].", ledger)

    assert "no verified company claim reference" in missing
    assert "company-disclosed fact lacks a verified C claim reference" in missing
    assert "unknown claim reference C99" in unknown


def test_evidence_ledger_filters_boilerplate_tables_and_incomplete_fragments() -> None:
    from app.services.evidence_ledger_service import EvidenceLedgerService

    content = "\n".join([
        "TABLE OF CONTENTS PART I Item 1. Business Item 1A. Risk Factors.",
        "6,134 7,154 6,587 11,541 $13,373 $13,050 Current tax provision 2025 2024 2023.",
        "Backed by a dealer network, the company provides products and services through three business segments",
        "58 CATERPILLAR 2025 FORM 10-K The following table is a reconciliation of income taxes.",
        "The company reported 2025 sales and revenues of $67.589 billion across its business segments.",
    ])

    sentences = EvidenceLedgerService._sentences(content)

    assert sentences == ["The company reported 2025 sales and revenues of $67.589 billion across its business segments."]


def test_evidence_ledger_extracts_exact_financial_metric_rows() -> None:
    from app.services.evidence_ledger_service import EvidenceLedgerService

    content = (
        "Financial Highlights (in millions, except per share data) 2025 2024 2023 "
        "Total net revenue $ 182,447 $ 177,556 $ 158,104 "
        "Total noninterest expense 95,640 91,797 87,172. "
        "Balance Sheet and Liquidity Cash and cash equivalents of $9.4 billion Total debt of $50.3 billion."
    )

    quotes = EvidenceLedgerService._metric_quotes(content)

    assert "Total net revenue $ 182,447 $ 177,556 $ 158,104" in quotes
    assert "Cash and cash equivalents of $9.4 billion" in quotes
    assert "Total debt of $50.3 billion" in quotes
    claims = EvidenceLedgerService._metric_claims(content)
    assert (
        "Total net revenue $ 182,447 $ 177,556 $ 158,104 (table unit: USD millions)",
        "Total net revenue $ 182,447 $ 177,556 $ 158,104",
    ) in claims

    walmart_rows = EvidenceLedgerService._metric_claims(
        "Fiscal Years Ended (Amounts in millions) 2026 2025 2024 "
        "Walmart U.S. Net sales $ 482,975 $ 462,415 $ 441,817 "
        "Operating income $ 29,825 $ 29,348 $ 27,012"
    )
    assert (
        "Walmart U.S. Net sales $ 482,975 $ 462,415 $ 441,817 (table unit: USD millions)",
        "Walmart U.S. Net sales $ 482,975 $ 462,415 $ 441,817",
    ) in walmart_rows
    assert EvidenceLedgerService._metric_claims(
        "Net income attributable to Example $ 21,893 $ 19,436 $ 15,511"
    ) == []


def test_evidence_ledger_does_not_split_common_financial_abbreviations() -> None:
    from app.services.evidence_ledger_service import EvidenceLedgerService

    content = "The index ranked the U.S. as sixth among manufacturing economies. J.P. Morgan serves institutional clients globally."

    assert EvidenceLedgerService._sentences(content) == [
        "The index ranked the U.S. as sixth among manufacturing economies.",
        "J.P. Morgan serves institutional clients globally.",
    ]


def test_evidence_ledger_anchors_only_high_confidence_fact_matches() -> None:
    from app.schemas.chat import EvidenceClaim
    from app.services.evidence_ledger_service import EvidenceLedgerService

    claims = [EvidenceClaim(
        claim_id="C1",
        claim="The company reported revenue of $120 million for the fiscal year.",
        source_filename="annual-report.pdf",
        document_role="company_disclosure",
        evidence_quote="The company reported revenue of $120 million for the fiscal year.",
        category="financial",
    )]

    anchored = EvidenceLedgerService.anchor_references(
        "Company-disclosed facts: Revenue was $120 million (annual-report.pdf).", claims
    )
    unrelated = EvidenceLedgerService.anchor_references(
        "Company-disclosed facts: The company has a durable competitive moat.", claims
    )

    assert anchored.endswith("[C1]")
    assert "[C1]" not in unrelated


def test_evidence_ledger_normalizes_reference_brackets() -> None:
    from app.services.evidence_ledger_service import EvidenceLedgerService

    assert EvidenceLedgerService.anchor_references("Fact 【C1】 and context (I2).", []) == "Fact [C1] and context [I2]."


def test_strategy_rag_prefers_financial_statement_chunks_over_cover_pages() -> None:
    from app.services.rag_service import RAGService

    cover = "Revenue growth highlights UNITED STATES SECURITIES AND EXCHANGE COMMISSION 2025."
    statement = (
        "Consolidated Statements of Income, dollars in millions, years ended 2025 and 2024. "
        "Total revenue 182447 177556; net income 57048 58471; operating income 64000 62000."
    )

    assert RAGService._term_chunk_score(statement, "revenue") > RAGService._term_chunk_score(cover, "revenue")


def test_strategy_gate_rejects_request_for_present_audited_statements() -> None:
    from app.services.chat_service import ChatService

    answer = "Request the audited financial statements. " + "x" * 900
    issues = ChatService._strategy_structure_issues(answer, ["[annual-report.pdf] FORM 10-K annual report"])

    assert "requests an annual filing already present in evidence" in issues


def test_strategy_answer_falls_back_when_evidence_control_does_not_pass() -> None:
    from app.services.chat_service import ChatService

    service = ChatService.__new__(ChatService)
    service.llm_service = SimpleNamespace(generate=lambda _prompt: '{"revised_answer":"Unsupported claim","removed_or_reframed_claims":[],"evidence_control_passed":false}')

    answer, passed, issues = service._ground_strategy_answer(
        SimpleNamespace(company_name="Acme", industry="SaaS", stage="Seed"),
        "investment strategy",
        ["[filing.txt] Revenue was 120."],
        ["Audited financial statements are missing"],
    )

    assert passed is False
    assert issues
    assert "Unsupported claim" not in answer
    assert "Audited financial statements are missing" in answer


def test_strategy_answer_recovers_with_deterministic_claim_plan() -> None:
    from app.schemas.chat import EvidenceClaim
    from app.services.chat_service import ChatService

    service = ChatService.__new__(ChatService)
    service.llm_service = SimpleNamespace(generate=lambda _prompt: '{"revised_answer":"bad","evidence_control_passed":false}')
    claims = [
        EvidenceClaim(
            claim_id="C1",
            claim="Revenue was $120 million for the fiscal year.",
            source_filename="annual-report.pdf",
            document_role="company_disclosure",
            evidence_quote="Revenue was $120 million for the fiscal year.",
            category="financial",
        ),
        EvidenceClaim(
            claim_id="I1",
            claim="Industry demand was volatile during the year.",
            source_filename="industry-report.pdf",
            document_role="industry_context",
            evidence_quote="Industry demand was volatile during the year.",
            category="market",
        ),
    ]

    answer, passed, issues = service._ground_strategy_answer(
        SimpleNamespace(company_name="Acme", industry="SaaS", stage="Seed"),
        "investment strategy",
        [
            "[annual-report.pdf | role=company_disclosure] Revenue was $120 million for the fiscal year.",
            "[industry-report.pdf | role=industry_context] Industry demand was volatile during the year.",
        ],
        ["Valuation inputs are missing"],
        claims,
    )

    assert passed is True
    assert "deterministic evidence plan" in issues[0]
    assert "[C1]" in answer
    assert "[I1]" in answer
    assert "## Post-Investment" in answer


def test_strategy_answer_requires_explicit_true_evidence_control() -> None:
    import json

    from app.services.chat_service import ChatService

    service = ChatService.__new__(ChatService)
    staged_answer = """
## IC Summary
Company-disclosed facts are drawn from [filing.txt]. Analyst inference remains conditional on verification.
## Pre-Investment
Company-disclosed fact: revenue is documented in [filing.txt]. Analyst inference: quality cannot assess without reconciliation. Verification action: reconcile audited revenue to the source ledger. IC gate: the committee sets its acceptance threshold after the mandate is documented. Cannot assess valuation.
## During-Investment
Company-disclosed fact: the filing identifies the business. Analyst inference: execution remains conditional. Verification action: verify the transaction field against [terms.txt]. IC gate: approve only after legal and financial fields reconcile. Cannot assess position size.
## Post-Investment
Company-disclosed fact: the filing defines reporting fields. Analyst inference: monitoring can use verified fields. Verification action: reconcile each reporting period to [filing.txt]. IC gate: escalation thresholds must be set by the committee. Cannot assess exit timing.
## Evidence Gaps
Obtain the named ledger fields and signed terms. Do not infer missing values.
""".strip()
    staged_answer += "\n" + "Verification remains tied to primary evidence and no threshold is invented. " * 8
    service.llm_service = SimpleNamespace(generate=lambda _prompt: json.dumps({
        "revised_answer": staged_answer,
        "removed_or_reframed_claims": [],
        "evidence_control_passed": True,
    }))

    answer, passed, issues = service._ground_strategy_answer(
        SimpleNamespace(company_name="Acme", industry="SaaS", stage="Seed"),
        "investment strategy",
        ["[filing.txt] Revenue was documented.", "[terms.txt] Transaction fields were documented."],
        [],
    )

    assert passed is True
    assert issues == []
    assert "IC Summary" in answer


def test_strategy_structure_gate_rejects_generic_short_answers() -> None:
    from app.services.chat_service import ChatService

    issues = ChatService._strategy_structure_issues("The evidence is insufficient. Request more documents.", ["[filing.txt] facts"])

    assert "pre-investment" in issues
    assert "IC gate" in issues
    assert "answer too short for a staged IC work plan" in issues


def test_strategy_answer_gets_one_repair_attempt(monkeypatch) -> None:
    from app.services.chat_service import ChatService

    responses = iter([
        '{"revised_answer":"draft","removed_or_reframed_claims":[],"evidence_control_passed":false}',
        '{"revised_answer":"repaired answer [filing.txt]","removed_or_reframed_claims":["draft"],"evidence_control_passed":true}',
        '{"revised_answer":"repaired answer [filing.txt]","unsupported_or_overreaching_claims":[],"evidence_control_passed":true}',
    ])
    service = ChatService.__new__(ChatService)
    service.llm_service = SimpleNamespace(generate=lambda _prompt: next(responses))
    monkeypatch.setattr(ChatService, "_strategy_structure_issues", staticmethod(lambda answer, _context: [] if "repaired" in answer else ["quality"]))

    answer, passed, issues = service._ground_strategy_answer(
        SimpleNamespace(company_name="Acme", industry="SaaS", stage="Seed"),
        "pre-investment strategy",
        ["[filing.txt] Evidence"],
        ["Customer evidence is missing"],
    )

    assert passed is True
    assert issues == []
    assert answer == "repaired answer [filing.txt]"


def test_financial_lifecycle_enforces_closing_gates_and_versions_opinions(api_client) -> None:
    token = api_client.post(
        "/api/auth/register",
        json={"email": "lifecycle@example.com", "username": "lifecycle", "password": "strong-password"},
    ).json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    project_id = api_client.post(
        "/api/projects", headers=headers,
        json={"name": "Lifecycle QA", "company_name": "Acme", "industry": "SaaS", "stage": "Growth"},
    ).json()["id"]
    upload = api_client.post(
        f"/api/projects/{project_id}/files/upload", headers=headers,
        files={"upload_file": ("signed-closing.txt", b"Signed closing evidence for Acme transaction", "text/plain")},
    )
    file_id = upload.json()["file"]["id"]

    rejected = api_client.put(
        f"/api/projects/{project_id}/lifecycle/transaction", headers=headers,
        json={
            "status": "closed", "approval_status": "pending",
            "decision_rationale": "The committee has not approved this transaction yet.",
            "conditions_precedent": [{"id": "license", "label": "License renewal", "status": "pending"}],
            "evidence_file_ids": [file_id],
        },
    )
    assert rejected.status_code == 400

    transaction = api_client.put(
        f"/api/projects/{project_id}/lifecycle/transaction", headers=headers,
        json={
            "transaction_type": "equity", "currency": "CNY", "committed_amount": "50000000",
            "entry_valuation": "300000000", "ownership_pct": "16.6667", "status": "closed",
            "approval_status": "approved",
            "decision_rationale": "The investment committee approved closing after legal and financial reconciliation.",
            "conditions_precedent": [{
                "id": "license", "label": "License renewal", "status": "satisfied", "evidence_file_id": file_id,
            }],
            "evidence_file_ids": [file_id],
        },
    )
    assert transaction.status_code == 200
    assert transaction.json()["status"] == "closed"
    assert api_client.get(f"/api/projects/{project_id}", headers=headers).json()["investment_status"] == "post_investment"

    metric = api_client.post(
        f"/api/projects/{project_id}/lifecycle/metrics", headers=headers,
        json={
            "code": "monthly_liquidity", "name": "月末可用流动性", "unit": "万元",
            "direction": "higher_better", "baseline_value": "3000", "target_value": "3000",
            "watch_threshold": "2500", "breach_threshold": "2000", "owner": "投后管理组",
            "source_description": "经财务负责人签署的月度资金报表",
        },
    )
    assert metric.status_code == 201
    observation = api_client.post(
        f"/api/projects/{project_id}/lifecycle/metrics/{metric.json()['id']}/observations", headers=headers,
        json={"period_end": "2026-06-30", "value": "1800", "source_file_id": file_id, "note": "月末函证值"},
    )
    assert observation.status_code == 201
    assert observation.json()["status"] == "high"

    summary = api_client.get(f"/api/projects/{project_id}/lifecycle", headers=headers)
    assert summary.status_code == 200
    assert summary.json()["risks"][0]["severity"] == "high"
    assert summary.json()["opinions"][0]["recommendation"] == "escalate"
    version = summary.json()["opinions"][0]["version"]
    refreshed = api_client.post(f"/api/projects/{project_id}/lifecycle/opinions/refresh", headers=headers)
    assert refreshed.json()["version"] == version


def test_configured_data_sources_are_tenant_scoped_and_scheduled(api_client, monkeypatch) -> None:
    from app.services.research_service import ResearchService

    monkeypatch.setattr(ResearchService, "_validate_configured_public_url", staticmethod(lambda _url: None))
    token = api_client.post(
        "/api/auth/register",
        json={"email": "sources@example.com", "username": "sources", "password": "strong-password"},
    ).json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    project_id = api_client.post(
        "/api/projects", headers=headers,
        json={"name": "Sources QA", "company_name": "Acme", "industry": "Banking", "stage": "Post"},
    ).json()["id"]
    source = api_client.post(
        f"/api/projects/{project_id}/lifecycle/data-sources", headers=headers,
        json={
            "name": "Regulatory filing feed", "source_type": "regulator", "category": "legal",
            "url": "https://regulator.example.com/acme/filing.pdf", "cadence_hours": 24,
        },
    )
    assert source.status_code == 201
    assert source.json()["status"] == "scheduled"
    assert source.json()["next_run_at"] is not None

    queued: list[tuple[str, str]] = []
    monkeypatch.setattr(
        tasks.ingest_data_source_subscription_task,
        "delay",
        lambda source_id, owner_id: queued.append((source_id, owner_id)),
    )
    schedule_result = tasks.schedule_due_data_sources_task.run()
    assert schedule_result == {"queued": 1}
    assert queued[0][0] == source.json()["id"]

    other_token = api_client.post(
        "/api/auth/register",
        json={"email": "source-other@example.com", "username": "source-other", "password": "strong-password"},
    ).json()["access_token"]
    assert api_client.get(
        f"/api/projects/{project_id}/lifecycle",
        headers={"Authorization": f"Bearer {other_token}"},
    ).status_code == 404


def test_configured_data_source_rejects_non_public_network_targets() -> None:
    from app.services.research_service import ResearchService

    with pytest.raises(ValueError, match="HTTPS"):
        ResearchService._validate_configured_public_url("http://example.com/report.pdf")
    with pytest.raises(ValueError, match="non-public"):
        ResearchService._validate_configured_public_url("https://127.0.0.1/report.pdf")


def test_fixed_professional_evaluation_suite_passes_and_rejects_unsafe_opinion() -> None:
    from app.evals.runner import run_suite

    report = run_suite()
    assert report["passed"] is True
    assert report["minimum_passing_score"] >= 80
    unsafe = next(item for item in report["results"] if item["case_id"] == "reject_invented_buy_call")
    assert unsafe["passed"] is False
    assert any("unsupported numeric" in item for item in unsafe["critical_issues"])
